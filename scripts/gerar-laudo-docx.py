"""
Gera a folha de Laudo Técnico da HB em .docx — preenchível no Word e no Google Docs.

Por que .docx e não PDF: o cliente precisa digitar nos campos. Por que tabelas e
não campos de formulário do Word: content controls são a forma "certa" no Word,
mas o Google Docs os descarta na importação. Célula de tabela vazia funciona
idêntica nos dois, e ainda deixa o Tab pular de campo em campo.

    python scripts/gerar-laudo-docx.py

Saída: docs/HB-Laudo-Tecnico.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# Tokens da papelaria HB — mesmos hex do design system e do arquivo Figma.
NAVY = "1A2A6C"
NAVY_SOFT = "4A6CB3"
BAR = "2B3A8F"
ACCENT = "C2185B"
ACCENT_STRONG = "E91E63"
WARN = "E05252"
SLATE = "5F7A7A"
INK = "1C1C1C"
ROW = "F1F1F1"
RULE = "D9D9D9"

# Arial em vez de Inter: é a única sans presente por padrão no Word e no Google
# Docs ao mesmo tempo. Fonte ausente vira substituição automática e quebra todo
# o ritmo da página — previsível vale mais que ideal aqui.
FONTE = "Arial"

LARGURA_UTIL = Mm(210 - 36)  # 210mm menos 18mm de margem de cada lado


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(f"w:{k}"), v)
    return e


def sombrear(celula, hex_cor):
    """Preenche o fundo de uma célula (zebra, bloco de campo aberto, cabeçalho)."""
    celula._tc.get_or_add_tcPr().append(
        _el("w:shd", val="clear", color="auto", fill=hex_cor)
    )


def bordas(celula, baixo=None, cima=None):
    """
    Aplica só bordas horizontais. Nenhuma borda vertical, nunca: grade fechada
    faz o olho parar em cada célula em vez de correr a linha.
    """
    tcPr = celula._tc.get_or_add_tcPr()
    tcb = OxmlElement("w:tcBorders")
    for lado, cor in (("top", cima), ("bottom", baixo)):
        b = _el(f"w:{lado}", val="single" if cor else "nil", sz="6", space="0")
        if cor:
            b.set(qn("w:color"), cor)
        tcb.append(b)
    for lado in ("left", "right"):
        tcb.append(_el(f"w:{lado}", val="nil", sz="0", space="0"))
    tcPr.append(tcb)


def altura_minima(linha, mm):
    trPr = linha._tr.get_or_add_trPr()
    trPr.append(_el("w:trHeight", val=str(int(mm * 56.7)), hRule="atLeast"))


def texto(par, conteudo, tamanho, cor, negrito=False, espacamento=1.35):
    run = par.add_run(conteudo)
    run.font.name = FONTE
    run.font.size = Pt(tamanho)
    run.font.bold = negrito
    run.font.color.rgb = RGBColor.from_string(cor)
    par.paragraph_format.line_spacing = espacamento
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    return run


def paragrafo(doc, conteudo, tamanho, cor, negrito=False, antes=0, depois=0):
    p = doc.add_paragraph()
    texto(p, conteudo, tamanho, cor, negrito)
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(depois)
    return p


def tabela_nua(doc, linhas, colunas, larguras=None):
    """Tabela sem estilo — todas as bordas são aplicadas célula a célula."""
    t = doc.add_table(rows=linhas, cols=colunas)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for linha in t.rows:
        for i, cel in enumerate(linha.cells):
            bordas(cel)
            cel.width = larguras[i] if larguras else Mm(174 / colunas)
    return t


# ---------------------------------------------------------------- documento

doc = Document()

sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.left_margin = sec.right_margin = Mm(18)
sec.top_margin = Mm(14)
sec.bottom_margin = Mm(16)
sec.header_distance = Mm(0)

estilo = doc.styles["Normal"]
estilo.font.name = FONTE
estilo.font.size = Pt(9)
estilo.font.color.rgb = RGBColor.from_string(INK)
estilo.paragraph_format.space_after = Pt(0)
estilo.paragraph_format.line_spacing = 1.35

# --- Faixa navy sangrando de borda a borda ---
# O Word não sangra além da margem no corpo. A saída é pôr a faixa no cabeçalho
# e puxá-la com recuo negativo até a borda física do papel.
cab = sec.header
cab.is_linked_to_previous = False
faixa = cab.add_table(rows=1, cols=1, width=Mm(210))
faixa.autofit = False
cel_faixa = faixa.rows[0].cells[0]
cel_faixa.width = Mm(210)
sombrear(cel_faixa, BAR)
bordas(cel_faixa)
altura_minima(faixa.rows[0], 6)
p_faixa = cel_faixa.paragraphs[0]
p_faixa.paragraph_format.left_indent = Mm(-18)
p_faixa.paragraph_format.space_after = Pt(0)
faixa._tbl.tblPr.append(_el("w:tblInd", w="-1020", type="dxa"))

# --- Cabeçalho da empresa ---
paragrafo(doc, "HB Comércio & Assistência", 13, NAVY_SOFT, antes=6)
for linha in (
    "Rua Raul Cardoso, Nº131",
    "Campos dos Goytacazes, Rio de Janeiro, 28027",
    "(22) 99861-6139",
):
    paragrafo(doc, linha, 7.5, SLATE)

# --- Título + status ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(0)
texto(p, "Laudo Técnico", 24, NAVY, negrito=True, espacamento=1.1)
texto(p, "   Laudo Nº ", 8, SLATE)
texto(p, "________", 8, SLATE)

paragrafo(doc, "Emitido em ____/____/________", 9.5, ACCENT, negrito=True, antes=2)
paragrafo(
    doc,
    "Este laudo descreve o estado do equipamento na data da análise. "
    "Alterações posteriores por terceiros invalidam o parecer.",
    6,
    WARN,
)


def secao(titulo):
    """Rótulo de seção com filete embaixo — sem fundo chapado, que gasta tinta."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    texto(p, titulo, 8.5, NAVY, negrito=True)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    pbdr.append(_el("w:bottom", val="single", sz="6", space="2", color=RULE))
    pPr.append(pbdr)


def grade(campos):
    """
    Grade rótulo/valor em 2 colunas. A célula de valor fica vazia com filete
    embaixo: dá pra digitar (Tab pula de campo em campo) e dá pra preencher à
    caneta depois de imprimir — que é como a bancada realmente usa.
    """
    metade = Mm(87)
    t = tabela_nua(doc, len(campos) // 2, 2, [metade, metade])
    for i, rotulo in enumerate(campos):
        cel = t.rows[i // 2].cells[i % 2]
        cel.paragraphs[0].text = ""
        texto(cel.paragraphs[0], rotulo, 7.5, NAVY_SOFT, negrito=True)
        valor = cel.add_paragraph()
        valor.paragraph_format.space_after = Pt(5)
        texto(valor, "", 9, INK)
        bordas(cel, baixo=RULE)
        altura_minima(t.rows[i // 2], 11)


def campo_aberto(linhas_texto):
    """Bloco cinza para texto livre. Sem borda tracejada — isso grita 'formulário'."""
    t = tabela_nua(doc, 1, 1, [LARGURA_UTIL])
    cel = t.rows[0].cells[0]
    sombrear(cel, ROW)
    cel.paragraphs[0].text = ""
    texto(cel.paragraphs[0], "", 9, INK)
    altura_minima(t.rows[0], linhas_texto * 5.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def tabela_itens(colunas, qtd_linhas):
    """Cabeçalho navy com filete abaixo, zebra nas linhas pares, sem grade."""
    largura = Mm(174 / len(colunas))
    t = tabela_nua(doc, qtd_linhas + 1, len(colunas), [largura] * len(colunas))
    for i, titulo in enumerate(colunas):
        cel = t.rows[0].cells[i]
        cel.paragraphs[0].text = ""
        texto(cel.paragraphs[0], titulo, 8.5, NAVY, negrito=True)
        bordas(cel, baixo=RULE)
    altura_minima(t.rows[0], 6)
    for r in range(1, qtd_linhas + 1):
        for cel in t.rows[r].cells:
            cel.paragraphs[0].text = ""
            texto(cel.paragraphs[0], "", 9, INK)
            if r % 2 == 0:
                sombrear(cel, ROW)
        altura_minima(t.rows[r], 7)


secao("Dados do cliente")
grade(["Nome", "Telefone", "Documento", "Data de entrada"])

secao("Identificação do equipamento")
grade(["Tipo", "Marca / Modelo", "Nº de série", "Acessórios recebidos"])

secao("Defeito relatado pelo cliente")
campo_aberto(3)

secao("Análise técnica")
campo_aberto(5)

secao("Testes realizados")
tabela_itens(["Teste", "Resultado", "Observação"], 4)

secao("Componentes avaliados")
tabela_itens(["Componente", "Estado", "Ação recomendada"], 4)

secao("Conclusão / Parecer")
paragrafo(doc, "REPARÁVEL", 13, ACCENT_STRONG, negrito=True, antes=2)
paragrafo(doc, "", 9, INK)

paragrafo(
    doc,
    "Garantia de 90 dias sobre o serviço executado, contados da data de retirada, "
    "limitada ao defeito descrito neste laudo. Não cobre danos por queda, oxidação "
    "ou intervenção de terceiros.",
    7.5,
    SLATE,
    antes=10,
)

# --- Assinaturas ---
assin = tabela_nua(doc, 1, 2, [Mm(80), Mm(80)])
for i, papel in enumerate(["Técnico responsável", "Cliente"]):
    cel = assin.rows[0].cells[i]
    cel.paragraphs[0].text = ""
    texto(cel.paragraphs[0], "", 9, INK)
    bordas(cel, baixo=RULE)
    altura_minima(assin.rows[0], 14)

# Os rótulos vão numa segunda tabela para ficarem SOB o filete, não dentro dele.
rotulos = tabela_nua(doc, 1, 2, [Mm(80), Mm(80)])
for i, papel in enumerate(["Técnico responsável", "Cliente"]):
    cel = rotulos.rows[0].cells[i]
    cel.paragraphs[0].text = ""
    texto(cel.paragraphs[0], papel, 6.5, SLATE)

saida = Path(__file__).resolve().parent.parent / "docs" / "HB-Laudo-Tecnico.docx"
saida.parent.mkdir(parents=True, exist_ok=True)
doc.save(saida)
print(f"Gerado: {saida}")
